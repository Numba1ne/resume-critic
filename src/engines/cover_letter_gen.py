"""
Cover Letter Generator - Generates tailored cover letters.
"""
from typing import Dict, Optional
from docx import Document
from src.utils.docx_handler import create_docx, save_docx


class CoverLetterGenerator:
    """Generates tailored cover letters following best practices."""
    
    def __init__(self, jd_analysis: Dict, cv_data: Dict, company_research: Optional[str] = None):
        """
        Initialize cover letter generator.
        
        Args:
            jd_analysis: JD analysis from parser
            cv_data: CV data (personal info, experience, etc.)
            company_research: Company-specific research/details
        """
        self.jd_analysis = jd_analysis
        self.cv_data = cv_data
        self.company_research = company_research
        self.cover_letter = {
            'paragraph_1': '',
            'paragraph_2': '',
            'paragraph_3': '',
            'paragraph_4': '',
            'paragraph_5': ''
        }
    
    def generate_paragraph_1_hook(self, company_name: str, job_title: str, 
                                  company_detail: Optional[str] = None) -> str:
        """
        Generate THE HOOK: Company-specific enthusiasm.
        
        Args:
            company_name: Company name
            job_title: Job title
            company_detail: Specific company detail (from research)
            
        Returns:
            Paragraph 1 text
        """
        if company_detail:
            hook = f"I'm excited to apply for the {job_title} position at {company_name}. {company_detail}"
        else:
            hook = f"I'm excited to apply for the {job_title} position at {company_name}. Your company's mission and values align with my professional goals."
        
        self.cover_letter['paragraph_1'] = hook
        return hook
    
    def generate_paragraph_2_technical_match(self) -> str:
        """
        Generate TECHNICAL MATCH: List relevant skills.
        
        Returns:
            Paragraph 2 text
        """
        technical_skills = self.jd_analysis.get('technical_skills', [])[:5]
        required_skills = self.jd_analysis.get('required_skills', [])[:3]
        
        # Combine and use exact terminology from JD
        skills_to_mention = list(set(technical_skills + required_skills))[:4]
        
        if len(skills_to_mention) >= 3:
            text = f"I'm proficient in {skills_to_mention[0]}, have strong {skills_to_mention[1]} skills, and experience with {skills_to_mention[2]}"
            if len(skills_to_mention) > 3:
                text += f" and {skills_to_mention[3]}"
            text += "."
        elif len(skills_to_mention) == 2:
            text = f"I'm proficient in {skills_to_mention[0]} and have strong {skills_to_mention[1]} skills."
        elif len(skills_to_mention) == 1:
            text = f"I'm proficient in {skills_to_mention[0]} and have relevant experience in data analytics."
        else:
            text = "I have strong technical skills and relevant experience in data analytics."
        
        self.cover_letter['paragraph_2'] = text
        return text
    
    def generate_paragraph_3_experience_story(self, achievement: Optional[str] = None) -> str:
        """
        Generate EXPERIENCE STORY: Brief example with numbers.
        
        Args:
            achievement: Specific achievement to highlight
            
        Returns:
            Paragraph 3 text
        """
        if achievement:
            # Use provided achievement
            text = f"In my previous role, {achievement}."
        else:
            # Generate generic but relevant story
            text = "In my previous role, I led data analysis projects that resulted in significant improvements. For example, I developed automated reporting solutions that reduced manual work by 30% and improved data accuracy."
        
        self.cover_letter['paragraph_3'] = text
        return text
    
    def generate_paragraph_4_why_this_role(self, role_aspect: Optional[str] = None) -> str:
        """
        Generate WHY THIS ROLE: Specific excitement.
        
        Args:
            role_aspect: Specific aspect from JD that interests you
            
        Returns:
            Paragraph 4 text
        """
        if role_aspect:
            text = f"I'm particularly drawn to this role because {role_aspect}. The opportunity to contribute to data-driven decision making aligns with my passion for analytics."
        else:
            # Use JD structure or responsibilities
            responsibilities = self.jd_analysis.get('structure', {}).get('responsibilities', [])
            if responsibilities:
                aspect = responsibilities[0][:100] + "..." if len(responsibilities[0]) > 100 else responsibilities[0]
                text = f"I'm particularly drawn to this role because of the opportunity to {aspect.lower()}. This aligns with my experience and interests."
            else:
                text = "I'm particularly drawn to this role because it offers the opportunity to work on challenging data problems and make a meaningful impact."
        
        self.cover_letter['paragraph_4'] = text
        return text
    
    def generate_paragraph_5_close(self, location_detail: Optional[str] = None,
                                   availability: Optional[str] = None) -> str:
        """
        Generate CLOSE: Logistics and interest.
        
        Args:
            location_detail: Location information
            availability: Availability information
            
        Returns:
            Paragraph 5 text
        """
        parts = []
        
        if location_detail:
            parts.append(f"I'm {location_detail}")
        
        if availability:
            parts.append(f"available {availability}")
        
        if parts:
            text = f"{', '.join(parts)}, and I'm excited to discuss how I can contribute to your team."
        else:
            text = "I'm excited to discuss how I can contribute to your team and would welcome the opportunity to speak with you further."
        
        self.cover_letter['paragraph_5'] = text
        return text
    
    def apply_tone_rules(self, text: str) -> str:
        """
        Ensure human, natural tone.
        
        Args:
            text: Text to check
            
        Returns:
            Text with tone improvements
        """
        # Check for overly formal phrases
        formal_phrases = {
            "I am writing to express my interest": "I'm excited to apply",
            "I would like to": "I'd like to",
            "I have the ability to": "I can",
            "I possess": "I have"
        }
        
        for formal, casual in formal_phrases.items():
            text = text.replace(formal, casual)
        
        return text
    
    def generate_complete_letter(self, company_name: str, job_title: str,
                                company_detail: Optional[str] = None,
                                achievement: Optional[str] = None,
                                role_aspect: Optional[str] = None,
                                location_detail: Optional[str] = None,
                                availability: Optional[str] = None) -> str:
        """
        Generate complete cover letter.
        
        Args:
            company_name: Company name
            job_title: Job title
            company_detail: Company-specific detail
            achievement: Achievement to highlight
            role_aspect: Role aspect that interests you
            location_detail: Location information
            availability: Availability information
            
        Returns:
            Complete cover letter text
        """
        # Generate all paragraphs
        p1 = self.generate_paragraph_1_hook(company_name, job_title, company_detail)
        p2 = self.generate_paragraph_2_technical_match()
        p3 = self.generate_paragraph_3_experience_story(achievement)
        p4 = self.generate_paragraph_4_why_this_role(role_aspect)
        p5 = self.generate_paragraph_5_close(location_detail, availability)
        
        # Combine paragraphs
        full_text = f"{p1}\n\n{p2}\n\n{p3}\n\n{p4}\n\n{p5}"
        
        # Apply tone rules
        full_text = self.apply_tone_rules(full_text)
        
        return full_text
    
    def generate_docx(self, output_path: str, company_name: str, job_title: str,
                     **kwargs) -> Dict:
        """
        Generate cover letter as DOCX file.
        
        Args:
            output_path: Path to save cover letter
            company_name: Company name
            job_title: Job title
            **kwargs: Additional parameters for letter generation
            
        Returns:
            Dict with file path and word count
        """
        # Generate letter text
        letter_text = self.generate_complete_letter(company_name, job_title, **kwargs)
        
        # Create DOCX
        doc = create_docx()
        
        # Add personal info if available
        personal_info = self.cv_data.get('personal_info', {})
        if personal_info:
            name = personal_info.get('name', '')
            if name:
                doc.add_paragraph(name)
            
            contact = []
            if personal_info.get('email'):
                contact.append(personal_info['email'])
            if personal_info.get('phone'):
                contact.append(personal_info['phone'])
            if contact:
                doc.add_paragraph(' | '.join(contact))
            
            doc.add_paragraph('')  # Blank line
        
        # Add date
        from datetime import datetime
        doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        doc.add_paragraph('')
        
        # Add paragraphs
        paragraphs = letter_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        # Save
        save_docx(doc, output_path)
        
        # Count words
        word_count = len(letter_text.split())
        
        return {
            'file_path': output_path,
            'word_count': word_count,
            'under_400_words': word_count < 400
        }
