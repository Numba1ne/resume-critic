"""
DOCX file handling utilities.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Optional
import os


def read_docx(file_path: str) -> Document:
    """
    Read a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Document: python-docx Document object
    """
    return Document(file_path)


def read_docx_from_bytes(file_bytes: bytes) -> Document:
    """
    Read DOCX from bytes (e.g., from Streamlit upload).
    
    Args:
        file_bytes: DOCX file as bytes
        
    Returns:
        Document: python-docx Document object
    """
    import io
    return Document(io.BytesIO(file_bytes))


def extract_docx_text(doc: Document) -> str:
    """
    Extract all text from a DOCX document.
    
    Args:
        doc: Document object
        
    Returns:
        str: All text content
    """
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def get_docx_paragraphs(doc: Document) -> List[str]:
    """
    Get all paragraph texts from DOCX.
    
    Args:
        doc: Document object
        
    Returns:
        List of paragraph texts
    """
    return [p.text for p in doc.paragraphs]


def get_docx_tables(doc: Document) -> List:
    """
    Get all tables from DOCX.
    
    Args:
        doc: Document object
        
    Returns:
        List of table objects
    """
    return doc.tables


def get_docx_sections(doc: Document) -> List:
    """
    Get all sections from DOCX.
    
    Args:
        doc: Document object
        
    Returns:
        List of section objects
    """
    return doc.sections


def save_docx(doc: Document, output_path: str):
    """
    Save DOCX document to file.
    
    Args:
        doc: Document object
        output_path: Path to save file
    """
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def create_docx() -> Document:
    """
    Create a new DOCX document.
    
    Returns:
        Document: New empty document
    """
    return Document()


def add_heading_to_docx(doc: Document, text: str, level: int = 1):
    """
    Add a heading to DOCX document.
    
    Args:
        doc: Document object
        text: Heading text
        level: Heading level (1-9)
    """
    doc.add_heading(text, level=level)


def add_paragraph_to_docx(doc: Document, text: str, style: Optional[str] = None):
    """
    Add a paragraph to DOCX document.
    
    Args:
        doc: Document object
        text: Paragraph text
        style: Paragraph style (e.g., 'List Bullet')
    """
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(text)


def find_section_in_docx(doc: Document, section_name: str) -> Optional[int]:
    """
    Find a section heading in DOCX and return its index.
    
    Args:
        doc: Document object
        section_name: Section name to find
        
    Returns:
        Index of paragraph containing section, or None
    """
    section_name_upper = section_name.upper()
    for i, paragraph in enumerate(doc.paragraphs):
        if section_name_upper in paragraph.text.upper():
            return i
    return None


def get_fonts_used(doc: Document) -> set:
    """
    Get all fonts used in document.
    
    Args:
        doc: Document object
        
    Returns:
        Set of font names
    """
    fonts = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                fonts.add(run.font.name)
    return fonts


def has_graphics(doc: Document) -> bool:
    """
    Check if document contains images/graphics.
    
    Args:
        doc: Document object
        
    Returns:
        True if graphics found
    """
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                return True
    except:
        pass
    return False


def has_tables(doc: Document) -> bool:
    """
    Check if document contains tables.
    
    Args:
        doc: Document object
        
    Returns:
        True if tables found
    """
    return len(doc.tables) > 0


def has_headers_footers(doc: Document) -> bool:
    """
    Check if document has content in headers/footers.
    
    Args:
        doc: Document object
        
    Returns:
        True if headers/footers have content
    """
    for section in doc.sections:
        # Check header
        if section.header.paragraphs:
            for para in section.header.paragraphs:
                if para.text.strip():
                    return True
        # Check footer
        if section.footer.paragraphs:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    return True
    return False
