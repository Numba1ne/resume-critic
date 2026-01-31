"""
ATS Compatibility Checker - Validates CV formatting against ATS requirements.
"""
import os
import yaml
from typing import Dict, List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt


class ATSCompatibilityChecker:
    """Check CV for ATS compatibility."""
    
    SCORING_WEIGHTS = {
        'file_format': 30,
        'layout': 25,
        'font': 10,
        'graphics': 20,
        'headers_footers': 15
    }
    
    def __init__(self, cv_path: str, rules_path: Optional[str] = None):
        """
        Initialize ATS checker.
        
        Args:
            cv_path: Path to CV file
            rules_path: Path to ATS rules YAML file
        """
        self.cv_path = cv_path
        self.issues = []
        self.score = 100
        
        # Load rules
        if rules_path is None:
            rules_path = os.path.join('data', 'ats_rules.yaml')
        
        try:
            with open(rules_path, 'r') as f:
                self.rules = yaml.safe_load(f)
        except FileNotFoundError:
            # Use default rules if file not found
            self.rules = {
                'formatting': {
                    'file_formats': {
                        'allowed': ['.docx'],
                        'prohibited': ['.pdf', '.doc']
                    },
                    'fonts': {
                        'allowed': ['Calibri', 'Arial', 'Times New Roman'],
                        'size_range': {'min': 10, 'max': 12}
                    }
                },
                'scoring': {
                    'file_format_weight': 30,
                    'layout_weight': 25,
                    'font_weight': 10,
                    'graphics_weight': 20,
                    'headers_weight': 15
                }
            }
        
        # Load document if DOCX
        _, ext = os.path.splitext(cv_path)
        if ext.lower() == '.docx':
            try:
                self.doc = Document(cv_path)
            except Exception as e:
                raise ValueError(f"Could not read DOCX file: {str(e)}")
        else:
            self.doc = None
    
    def check_file_format(self) -> int:
        """
        Check if file is .docx.
        
        Returns:
            Deduction points (negative) or 0
        """
        _, ext = os.path.splitext(self.cv_path)
        allowed = self.rules['formatting']['file_formats']['allowed']
        
        if ext.lower() not in allowed:
            deduction = -self.SCORING_WEIGHTS['file_format']
            self.issues.append({
                'severity': 'HIGH',
                'issue': f'File format is {ext}, should be {", ".join(allowed)}',
                'deduction': abs(deduction)
            })
            return deduction
        
        return 0
    
    def check_layout(self) -> int:
        """
        Check for tables, columns, text boxes.
        
        Returns:
            Deduction points (negative) or 0
        """
        if not self.doc:
            return 0
        
        deduction = 0
        
        # Check for tables
        if len(self.doc.tables) > 0:
            self.issues.append({
                'severity': 'HIGH',
                'issue': f'Found {len(self.doc.tables)} table(s) - ATS may not parse correctly',
                'deduction': 15
            })
            deduction += 15
        
        # Check for multiple columns (simplified check)
        for section in self.doc.sections:
            if hasattr(section, 'columns'):
                # Check if more than one column
                try:
                    if section.columns.count > 1:
                        self.issues.append({
                            'severity': 'HIGH',
                            'issue': 'Multi-column layout detected',
                            'deduction': 10
                        })
                        deduction += 10
                        break
                except:
                    pass
        
        return -deduction
    
    def check_fonts(self) -> int:
        """
        Check for standard fonts.
        
        Returns:
            Deduction points (negative) or 0
        """
        if not self.doc:
            return 0
        
        allowed_fonts = self.rules['formatting']['fonts']['allowed']
        font_size_range = self.rules['formatting']['fonts']['size_range']
        min_size = font_size_range.get('min', 10)
        max_size = font_size_range.get('max', 12)
        
        deduction = 0
        fonts_used = set()
        non_standard_fonts = set()
        invalid_sizes = []
        
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                    if run.font.name not in allowed_fonts:
                        non_standard_fonts.add(run.font.name)
                
                if run.font.size:
                    size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else None
                    if size_pt and (size_pt < min_size or size_pt > max_size):
                        invalid_sizes.append(size_pt)
        
        if non_standard_fonts:
            self.issues.append({
                'severity': 'MEDIUM',
                'issue': f'Non-standard fonts detected: {", ".join(non_standard_fonts)}',
                'deduction': self.SCORING_WEIGHTS['font']
            })
            deduction += self.SCORING_WEIGHTS['font']
        
        if invalid_sizes:
            unique_sizes = list(set(invalid_sizes))
            self.issues.append({
                'severity': 'LOW',
                'issue': f'Font sizes outside recommended range ({min_size}-{max_size}pt): {", ".join(map(str, unique_sizes))}',
                'deduction': 5
            })
            deduction += 5
        
        return -deduction
    
    def check_headers_footers(self) -> int:
        """
        Check if headers/footers are used.
        
        Returns:
            Deduction points (negative) or 0
        """
        if not self.doc:
            return 0
        
        deduction = 0
        
        for section in self.doc.sections:
            # Check header
            if section.header.paragraphs:
                has_content = any(p.text.strip() for p in section.header.paragraphs)
                if has_content:
                    self.issues.append({
                        'severity': 'MEDIUM',
                        'issue': 'Header contains text - ATS may not read it',
                        'deduction': 8
                    })
                    deduction += 8
            
            # Check footer
            if section.footer.paragraphs:
                has_content = any(p.text.strip() for p in section.footer.paragraphs)
                if has_content:
                    self.issues.append({
                        'severity': 'MEDIUM',
                        'issue': 'Footer contains text - ATS may not read it',
                        'deduction': 7
                    })
                    deduction += 7
        
        return -deduction
    
    def check_graphics(self) -> int:
        """
        Check for images, shapes, etc.
        
        Returns:
            Deduction points (negative) or 0
        """
        if not self.doc:
            return 0
        
        deduction = 0
        
        # Check for inline shapes (images, etc.)
        has_graphics = False
        try:
            for rel in self.doc.part.rels.values():
                if "image" in rel.target_ref.lower():
                    has_graphics = True
                    break
        except:
            pass
        
        if has_graphics:
            self.issues.append({
                'severity': 'HIGH',
                'issue': 'Images/graphics detected - ATS cannot read these',
                'deduction': self.SCORING_WEIGHTS['graphics']
            })
            deduction += self.SCORING_WEIGHTS['graphics']
        
        return -deduction
    
    def check_special_characters(self) -> int:
        """
        Check for unusual bullet points or special characters.
        
        Returns:
            Deduction points (negative) or 0
        """
        if not self.doc:
            return 0
        
        # Extract all text
        text = '\n'.join([p.text for p in self.doc.paragraphs])
        
        # Standard bullet is • or -
        unusual_bullets = ['→', '►', '▪', '▫', '■', '□', '◆', '◇']
        found_unusual = [b for b in unusual_bullets if b in text]
        
        if found_unusual:
            self.issues.append({
                'severity': 'LOW',
                'issue': f'Unusual bullet points detected: {", ".join(found_unusual)}',
                'deduction': 5
            })
            return -5
        
        return 0
    
    def calculate_total_score(self) -> int:
        """
        Run all checks and calculate final score.
        
        Returns:
            Score from 0-100
        """
        self.score = 100
        
        self.score += self.check_file_format()
        self.score += self.check_layout()
        self.score += self.check_fonts()
        self.score += self.check_headers_footers()
        self.score += self.check_graphics()
        self.score += self.check_special_characters()
        
        return max(0, self.score)  # Can't go below 0
    
    def generate_report(self) -> Dict:
        """
        Generate detailed report.
        
        Returns:
            Dict with score, grade, issues, and recommendations
        """
        total_score = self.calculate_total_score()
        
        report = {
            'score': total_score,
            'grade': self._get_grade(total_score),
            'issues': sorted(self.issues, key=lambda x: x['deduction'], reverse=True),
            'recommendations': self._get_recommendations()
        }
        
        return report
    
    def _get_grade(self, score: int) -> str:
        """
        Convert score to letter grade.
        
        Args:
            score: ATS compatibility score
            
        Returns:
            Letter grade with description
        """
        if score >= 90:
            return 'A - Excellent'
        elif score >= 80:
            return 'B - Good'
        elif score >= 70:
            return 'C - Acceptable'
        elif score >= 60:
            return 'D - Needs Improvement'
        else:
            return 'F - Major Issues'
    
    def _get_recommendations(self) -> List[str]:
        """
        Generate fix recommendations.
        
        Returns:
            List of recommendation strings
        """
        recommendations = []
        
        for issue in self.issues:
            issue_lower = issue['issue'].lower()
            
            if 'table' in issue_lower:
                recommendations.append('Remove all tables and use plain text with bullets')
            elif 'column' in issue_lower:
                recommendations.append('Convert to single-column layout')
            elif 'font' in issue_lower:
                recommendations.append('Change all fonts to Calibri, Arial, or Times New Roman (10-12pt)')
            elif 'header' in issue_lower or 'footer' in issue_lower:
                recommendations.append('Move all header/footer content into main document body')
            elif 'image' in issue_lower or 'graphic' in issue_lower:
                recommendations.append('Remove all images, logos, and graphics')
            elif 'format' in issue_lower:
                recommendations.append('Convert file to .docx format')
        
        return list(set(recommendations))  # Remove duplicates
