"""
Quality validator for CVs - Checks for best practices.
"""
import os
import re
from typing import Dict, List
from docx import Document


class QualityValidator:
    """Validate CV quality and best practices."""
    
    def __init__(self, cv_path: str):
        """
        Initialize quality validator.
        
        Args:
            cv_path: Path to CV file
        """
        self.cv_path = cv_path
        self.warnings = []
        
        # Load document
        _, ext = os.path.splitext(cv_path)
        if ext.lower() == '.docx':
            self.doc = Document(cv_path)
            self.cv_text = '\n'.join([p.text for p in self.doc.paragraphs])
        else:
            self.doc = None
            with open(cv_path, 'r', encoding='utf-8', errors='ignore') as f:
                self.cv_text = f.read()
        
        self.cv_text_lower = self.cv_text.lower()
    
    def check_quantifiable_achievements(self) -> List[str]:
        """
        Check if achievements have quantifiable metrics.
        
        Returns:
            List of warnings for achievements without numbers
        """
        warnings = []
        
        # Find bullet points
        bullets = re.findall(r'[•\-\*]\s*(.+?)(?=\n|$)', self.cv_text, re.MULTILINE)
        
        # Check for numbers in bullets
        number_pattern = r'\d+[%£$KM]|\d+%|\d+\s*(?:percent|million|thousand|k|m)'
        
        for bullet in bullets:
            if not re.search(number_pattern, bullet, re.IGNORECASE):
                # Check if it's an achievement bullet (starts with action verb)
                action_verbs = ['led', 'managed', 'increased', 'decreased', 'improved', 
                               'reduced', 'achieved', 'delivered', 'created', 'built']
                if any(bullet.lower().startswith(verb) for verb in action_verbs):
                    warnings.append(f"Consider adding quantifiable metrics: {bullet[:50]}...")
        
        return warnings[:5]  # Limit to 5 warnings
    
    def check_job_title_match(self, target_title: str) -> bool:
        """
        Check if CV job title matches target job title.
        
        Args:
            target_title: Target job title from JD
            
        Returns:
            True if match found
        """
        # Extract CV title (usually first large text)
        cv_title = self._extract_cv_title()
        
        # Check similarity
        target_words = set(target_title.lower().split())
        cv_words = set(cv_title.lower().split())
        
        # Calculate overlap
        overlap = len(target_words & cv_words)
        similarity = overlap / len(target_words) if target_words else 0
        
        return similarity >= 0.5  # 50% word overlap
    
    def _extract_cv_title(self) -> str:
        """Extract job title from CV."""
        if self.doc:
            # Check first few paragraphs for large/bold text
            for para in self.doc.paragraphs[:5]:
                if para.runs:
                    first_run = para.runs[0]
                    if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                        return para.text.strip()
        
        # Fallback: first line
        return self.cv_text.split('\n')[0].strip()
    
    def check_keyword_density(self, required_keywords: List[str]) -> Dict:
        """
        Check if required keywords are present.
        
        Args:
            required_keywords: List of required keywords
            
        Returns:
            Dict with keyword presence info
        """
        present = []
        missing = []
        
        for keyword in required_keywords:
            pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
            if re.search(pattern, self.cv_text_lower):
                present.append(keyword)
            else:
                missing.append(keyword)
        
        return {
            'present': present,
            'missing': missing,
            'coverage': len(present) / len(required_keywords) * 100 if required_keywords else 0
        }
    
    def validate_all(self, target_title: str = None, required_keywords: List[str] = None) -> Dict:
        """
        Run all quality checks.
        
        Args:
            target_title: Target job title
            required_keywords: Required keywords
            
        Returns:
            Dict with all validation results
        """
        results = {
            'quantifiable_warnings': self.check_quantifiable_achievements(),
            'title_match': None,
            'keyword_coverage': None
        }
        
        if target_title:
            results['title_match'] = self.check_job_title_match(target_title)
        
        if required_keywords:
            results['keyword_coverage'] = self.check_keyword_density(required_keywords)
        
        return results
