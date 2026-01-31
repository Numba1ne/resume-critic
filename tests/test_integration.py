"""
Integration tests for end-to-end workflows.
"""
import pytest
import os
from pathlib import Path
from docx import Document
from src.parsers.jd_parser import JobDescriptionParser
from src.engines.cv_tailor import CVTailoringEngine
from src.checkers.ats_compatibility import ATSCompatibilityChecker
from src.utils.database import Database


def create_test_cv(file_path: str):
    """Create a test CV DOCX."""
    doc = Document()
    doc.add_heading('Data Analyst', 0)
    doc.add_paragraph('John Doe')
    doc.add_paragraph('john.doe@email.com')
    doc.add_heading('PROFESSIONAL EXPERIENCE', 1)
    doc.add_paragraph('• Led data analysis projects')
    doc.add_heading('SKILLS', 1)
    doc.add_paragraph('SQL, Python, Excel')
    doc.save(file_path)


def test_end_to_end_workflow(tmp_path):
    """Test complete workflow from JD to tailored CV."""
    # Create test CV
    cv_path = tmp_path / "base_cv.docx"
    create_test_cv(str(cv_path))
    
    # Parse JD
    jd_text = """
    Senior Data Analyst Position
    
    Required:
    - Strong SQL skills
    - Python programming
    - Tableau experience
    
    We value: Innovation, Collaboration
    """
    parser = JobDescriptionParser(jd_text)
    analysis = parser.extract_all()
    
    # Generate tailored CV
    output_path = tmp_path / "tailored_cv.docx"
    tailor = CVTailoringEngine(str(cv_path), analysis)
    result = tailor.generate_tailored_cv(str(output_path))
    
    assert result['success']
    assert os.path.exists(output_path)
    
    # Check ATS compatibility
    checker = ATSCompatibilityChecker(str(output_path))
    ats_report = checker.generate_report()
    
    assert ats_report['score'] >= 0
    assert 'grade' in ats_report


def test_database_operations(tmp_path):
    """Test database operations."""
    db_path = tmp_path / "test.db"
    db = Database(str(db_path))
    
    # Add application
    app_id = db.add_application("Test Company", "Test Role")
    assert app_id > 0
    
    # Get application
    app = db.get_application(app_id)
    assert app is not None
    assert app['company_name'] == "Test Company"
    
    # Update status
    db.update_application(app_id, status="Interview")
    app = db.get_application(app_id)
    assert app['status'] == "Interview"
