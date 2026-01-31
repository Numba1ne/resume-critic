"""
Tests for ATS Compatibility Checker.
"""
import pytest
import os
from pathlib import Path
from docx import Document
from src.checkers.ats_compatibility import ATSCompatibilityChecker


def create_test_docx(file_path: str):
    """Create a test DOCX file."""
    doc = Document()
    doc.add_heading('Test CV', 0)
    doc.add_paragraph('This is a test CV for ATS checking.')
    doc.save(file_path)


def test_file_format_check(tmp_path):
    """Test file format checking."""
    test_file = tmp_path / "test.docx"
    create_test_docx(str(test_file))
    
    checker = ATSCompatibilityChecker(str(test_file))
    score = checker.check_file_format()
    assert score == 0  # No deduction for .docx


def test_ats_scoring(tmp_path):
    """Test ATS scoring."""
    test_file = tmp_path / "test.docx"
    create_test_docx(str(test_file))
    
    checker = ATSCompatibilityChecker(str(test_file))
    total_score = checker.calculate_total_score()
    assert 0 <= total_score <= 100


def test_report_generation(tmp_path):
    """Test report generation."""
    test_file = tmp_path / "test.docx"
    create_test_docx(str(test_file))
    
    checker = ATSCompatibilityChecker(str(test_file))
    report = checker.generate_report()
    
    assert 'score' in report
    assert 'grade' in report
    assert 'issues' in report
    assert 'recommendations' in report
